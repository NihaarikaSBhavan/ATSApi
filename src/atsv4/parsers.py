from __future__ import annotations

from pathlib import Path

from .models import ParsedDocument
from .sectioning import split_into_sections


def parse_document(path: str | Path) -> ParsedDocument:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        raw_text, metadata = _parse_pdf(path)
    elif suffix == ".docx":
        raw_text, metadata = _parse_docx(path)
    else:
        raw_text = path.read_text(encoding="utf-8")
        metadata = {"parser": "plain_text"}

    return build_parsed_document(
        source_path=str(path),
        raw_text=raw_text,
        metadata=metadata,
    )


def parse_text_document(text: str, source_name: str = "inline_text") -> ParsedDocument:
    return build_parsed_document(
        source_path=source_name,
        raw_text=text,
        metadata={"parser": "inline_text"},
    )


def build_parsed_document(
    source_path: str,
    raw_text: str,
    metadata: dict[str, str] | None = None,
) -> ParsedDocument:
    normalized_text = normalize_whitespace(raw_text)
    sections = split_into_sections(normalized_text)
    return ParsedDocument(
        source_path=source_path,
        raw_text=normalized_text,
        sections=sections,
        metadata=metadata or {},
    )


def normalize_whitespace(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    cleaned = [line for line in lines if line]
    return "\n".join(cleaned)


def _parse_pdf(path: Path) -> tuple[str, dict[str, str]]:
    errors: list[str] = []

    try:
        import fitz  # type: ignore

        doc = fitz.open(path)
        text = "\n".join(page.get_text("text") for page in doc)
        if text.strip():
            return text, {"parser": "pymupdf"}
    except Exception as exc:  # pragma: no cover
        errors.append(f"pymupdf: {exc}")

    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(path) as pdf:
            text = "\n".join((page.extract_text(x_tolerance=2, y_tolerance=3) or "") for page in pdf.pages)
        if text.strip():
            return text, {"parser": "pdfplumber"}
    except Exception as exc:  # pragma: no cover
        errors.append(f"pdfplumber: {exc}")

    error_summary = "; ".join(errors) if errors else "no parser available"
    raise RuntimeError(f"Unable to parse PDF '{path}': {error_summary}")


def _parse_docx(path: Path) -> tuple[str, dict[str, str]]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required to parse DOCX files") from exc

    doc = Document(path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    return text, {"parser": "python-docx"}
